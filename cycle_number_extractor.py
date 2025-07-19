#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
循环数提取器
从"零数的使用.docx"文档中提取循环数相关信息
"""

import json
import re
from docx import Document
from typing import Dict, List, Any

class CycleNumberExtractor:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.cycle_numbers = {}
        
    def extract_cycle_numbers(self) -> Dict[str, Any]:
        """提取所有循环数信息"""
        print("开始提取循环数信息...")
        
        # 提取循环数1-9
        for i in range(1, 10):
            self._extract_cycle_number(str(i))
            
        # 提取特殊循环数
        special_numbers = ['11', '22', '33']
        for num in special_numbers:
            self._extract_cycle_number(num)
            
        # 提取人生周期表信息
        self._extract_life_cycle_tables()
        
        return self.cycle_numbers
    
    def _extract_cycle_number(self, number: str):
        """提取特定循环数的信息"""
        print(f"提取循环数{number}...")
        
        cycle_info = {
            "number": number,
            "meaning": "",
            "description": "",
            "key_points": [],
            "life_aspects": {
                "love_marriage": "",
                "work_finance": "",
                "health_life": ""
            }
        }
        
        # 查找循环数标题
        title_pattern = f'循环数{number}'
        content = []
        in_section = False
        
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # 检查是否是循环数标题
            if title_pattern in text and "的意义" in text:
                in_section = True
                continue
                
            # 如果遇到下一个循环数标题，结束当前提取
            if in_section and "循环数" in text and "的意义" in text and title_pattern not in text:
                break
                
            if in_section:
                content.append(text)
        
        if content:
            # 处理提取的内容
            full_text = " ".join(content)
            
            # 提取意义
            meaning_match = re.search(r'代表「([^」]+)」', full_text)
            if meaning_match:
                cycle_info["meaning"] = meaning_match.group(1)
            
            # 提取描述
            description_parts = []
            for line in content:
                if "代表" in line or "象征" in line or "比喻" in line:
                    description_parts.append(line)
            
            if description_parts:
                cycle_info["description"] = " ".join(description_parts[:3])  # 取前3个描述
            
            # 提取关键要点
            key_points = []
            for line in content:
                if any(keyword in line for keyword in ["重点", "关键", "重要", "注意", "建议"]):
                    key_points.append(line)
            
            cycle_info["key_points"] = key_points[:5]  # 最多5个要点
            
            # 提取生活方面信息
            self._extract_life_aspects(content, cycle_info)
            
            self.cycle_numbers[f"cycle_{number}"] = cycle_info
    
    def _extract_life_aspects(self, content: List[str], cycle_info: Dict):
        """提取生活方面的信息"""
        current_aspect = None
        
        for line in content:
            if "恋爱" in line or "婚姻" in line:
                current_aspect = "love_marriage"
                continue
            elif "工作" in line or "财运" in line:
                current_aspect = "work_finance"
                continue
            elif "健康" in line or "生活" in line:
                current_aspect = "health_life"
                continue
            
            if current_aspect and line and len(line) > 10:
                if not cycle_info["life_aspects"][current_aspect]:
                    cycle_info["life_aspects"][current_aspect] = line
                else:
                    cycle_info["life_aspects"][current_aspect] += " " + line
    
    def _extract_life_cycle_tables(self):
        """提取人生周期表信息"""
        print("提取人生周期表...")
        
        tables = []
        in_table_section = False
        
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            if "人生周期" in text and "生命阶段" in text:
                in_table_section = True
                continue
                
            if in_table_section and text and any(char.isdigit() for char in text):
                # 检查是否是表格行
                if re.search(r'\d+\s+\d+', text):
                    tables.append(text)
        
        if tables:
            self.cycle_numbers["life_cycle_tables"] = {
                "description": "人生周期表 - 显示不同命运数的人生阶段",
                "tables": tables[:10]  # 最多10个表格
            }
    
    def save_to_json(self, output_path: str):
        """保存到JSON文件"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.cycle_numbers, f, ensure_ascii=False, indent=2)
        print(f"已保存到: {output_path}")
    
    def generate_statistics(self) -> str:
        """生成统计信息"""
        stats = []
        stats.append("循环数提取统计")
        stats.append("=" * 50)
        
        total_cycles = len([k for k in self.cycle_numbers.keys() if k.startswith('cycle_')])
        stats.append(f"提取的循环数数量: {total_cycles}")
        
        for key, value in self.cycle_numbers.items():
            if key.startswith('cycle_'):
                stats.append(f"\n循环数{value['number']}:")
                stats.append(f"  意义: {value['meaning']}")
                stats.append(f"  关键要点数: {len(value['key_points'])}")
                stats.append(f"  生活方面: {sum(1 for v in value['life_aspects'].values() if v)}")
        
        if "life_cycle_tables" in self.cycle_numbers:
            stats.append(f"\n人生周期表: {len(self.cycle_numbers['life_cycle_tables']['tables'])} 个表格")
        
        return "\n".join(stats)

def main():
    """主函数"""
    extractor = CycleNumberExtractor("零数的使用.docx")
    
    # 提取数据
    cycle_data = extractor.extract_cycle_numbers()
    
    # 保存JSON文件
    extractor.save_to_json("cycle_numbers.json")
    
    # 生成统计信息
    stats = extractor.generate_statistics()
    print("\n" + stats)
    
    # 保存统计信息
    with open("cycle_statistics.txt", "w", encoding="utf-8") as f:
        f.write(stats)
    
    print(f"\n提取完成！")
    print(f"- JSON文件: cycle_numbers.json")
    print(f"- 统计文件: cycle_statistics.txt")

if __name__ == "__main__":
    main() 