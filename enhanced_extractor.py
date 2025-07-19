#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
日期匹配信息提取工具 - 增强版
实现双向匹配关系：如果A在B的X分类下，那么B也应该在A的X分类下

功能特点：
1. 自动识别主日期标题（包含"相关生日如下"）
2. 正确识别四种分类：情人伴侣、工作伙伴朋友、竞争对手天敌、灵魂伴侣
3. 提取每个分类下的所有匹配日期
4. 自动补充缺失的双向匹配关系
5. 生成JSON、Markdown、CSV等多种格式输出
6. 生成反向索引，支持按日期查询
7. 提供详细的统计信息

使用方法：
python3 enhanced_extractor.py

输出文件：
- enhanced_date_matches.json: 增强版JSON数据
- enhanced_reverse_index.json: 增强版反向索引
- enhanced_date_matches.md: 增强版Markdown表格
- enhanced_date_matches.csv: 增强版CSV文件
- enhanced_statistics.txt: 增强版统计信息
"""

import json
import re
import csv
from docx import Document
from typing import Dict, List, Optional, Set

class EnhancedDateExtractor:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        # 定义分类标题的各种可能格式
        self.category_patterns = {
            "情人伴侣": [r"情人\s*伴侣", r"情人伴侣"],
            "工作伙伴朋友": [r"工作伙伴\s*朋友", r"工作伙伴朋友"],
            "竞争对手天敌": [r"竞争对手\s*天敌", r"竞争对手天敌"],
            "灵魂伴侣": [r"灵魂伴侣"]
        }
        
    def clean_text(self, text: str) -> str:
        """清理文本，去除特殊字符和空格"""
        return text.replace("　", "").replace("\xa0", "").replace(" ", "")
        
    def is_category_title(self, text: str) -> Optional[str]:
        """检查文本是否为分类标题"""
        cleaned_text = self.clean_text(text)
        for category, patterns in self.category_patterns.items():
            for pattern in patterns:
                if re.match(pattern, cleaned_text):
                    return category
        return None
    
    def is_main_date_title(self, text: str) -> Optional[str]:
        """检查文本是否为主日期标题（包含"相关生日如下"）"""
        cleaned_text = self.clean_text(text)
        match = re.match(r'^(\d{1,2}月\d{1,2}日)\s*相关生日如下[：:]*', cleaned_text)
        if match:
            return match.group(1)
        return None
    
    def extract_dates(self) -> List[Dict]:
        """提取所有日期匹配信息"""
        results = []
        current_main_date = None
        current_matches = {}
        current_category = None
        
        i = 0
        while i < len(self.doc.paragraphs):
            text = self.doc.paragraphs[i].text.strip()
            if not text:
                i += 1
                continue
                
            # 检查是否为主日期标题
            main_date = self.is_main_date_title(text)
            if main_date:
                # 保存前一个主日期的数据
                if current_main_date and current_matches:
                    results.append({
                        "主日期": current_main_date,
                        "匹配": current_matches.copy()
                    })
                
                # 开始新的主日期
                current_main_date = main_date
                current_matches = {cat: [] for cat in self.category_patterns.keys()}
                current_category = None
                i += 1
                continue
            
            # 检查是否为分类标题
            category_found = self.is_category_title(text)
            if category_found and current_main_date:
                current_category = category_found
                i += 1
                continue
            
            # 如果当前有主日期和分类，则提取日期
            if current_main_date and current_category:
                dates = self.extract_dates_from_text(text)
                if dates:
                    if current_category not in current_matches:
                        current_matches[current_category] = []
                    current_matches[current_category].extend(dates)
            else:
                print(f"[跳过段落] 未知分类或主日期丢失：'{text}'")
            
            i += 1
        
        # 添加最后一个主日期的数据
        if current_main_date and current_matches:
            results.append({
                "主日期": current_main_date,
                "匹配": current_matches.copy()
            })
        
        return results
    
    def extract_dates_from_text(self, text: str) -> List[str]:
        """从文本中提取日期"""
        # 匹配 "X月X日" 格式的日期
        date_pattern = r'\d{1,2}月\d{1,2}日'
        dates = re.findall(date_pattern, text)
        
        # 清理和验证日期
        cleaned_dates = []
        for date in dates:
            # 验证日期格式
            if self.is_valid_date(date):
                cleaned_dates.append(date)
            else:
                print(f"警告: 发现异常日期格式: {date}")
        
        return cleaned_dates
    
    def is_valid_date(self, date_str: str) -> bool:
        """验证日期格式是否有效"""
        match = re.match(r'(\d{1,2})月(\d{1,2})日', date_str)
        if not match:
            return False
        
        month, day = int(match.group(1)), int(match.group(2))
        return 1 <= month <= 12 and 1 <= day <= 31
    
    def enhance_bidirectional_matches(self, data: List[Dict]) -> List[Dict]:
        """增强双向匹配关系"""
        print("正在增强双向匹配关系...")
        
        # 创建主日期到数据的映射
        main_date_map = {item["主日期"]: item for item in data}
        
        # 统计增强前的匹配数
        original_matches = sum(
            sum(len(dates) for dates in item["匹配"].values())
            for item in data
        )
        
        # 收集所有需要添加的匹配关系
        additions = []
        
        for item in data:
            main_date = item["主日期"]
            matches = item["匹配"]
            
            for category, dates in matches.items():
                for date in dates:
                    # 检查反向匹配是否存在
                    if date in main_date_map:
                        target_item = main_date_map[date]
                        if category not in target_item["匹配"]:
                            target_item["匹配"][category] = []
                        
                        # 如果反向匹配不存在，则添加
                        if main_date not in target_item["匹配"][category]:
                            target_item["匹配"][category].append(main_date)
                            additions.append({
                                "from": main_date,
                                "to": date,
                                "category": category
                            })
        
        # 统计增强后的匹配数
        enhanced_matches = sum(
            sum(len(dates) for dates in item["匹配"].values())
            for item in data
        )
        
        print(f"增强前总匹配数: {original_matches}")
        print(f"增强后总匹配数: {enhanced_matches}")
        print(f"新增匹配数: {enhanced_matches - original_matches}")
        print(f"新增匹配关系: {len(additions)} 个")
        
        return data
    
    def generate_reverse_index(self, data: List[Dict]) -> Dict[str, List[Dict]]:
        """生成反向索引：输入任意日期，返回它在哪些主日期中出现过"""
        reverse_index = {}
        
        for item in data:
            main_date = item["主日期"]
            matches = item["匹配"]
            
            for category, dates in matches.items():
                for date in dates:
                    if date not in reverse_index:
                        reverse_index[date] = []
                    reverse_index[date].append({
                        "主日期": main_date,
                        "分类": category
                    })
        
        return reverse_index
    
    def export_to_markdown(self, data: List[Dict], output_path: str):
        """导出为Markdown表格格式"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# 增强版日期匹配表\n\n")
            f.write("| 主日期 | 情人伴侣 | 工作伙伴朋友 | 竞争对手天敌 | 灵魂伴侣 |\n")
            f.write("|--------|----------|--------------|--------------|----------|\n")
            
            for item in data:
                main_date = item["主日期"]
                matches = item["匹配"]
                
                row = [main_date]
                for category in self.category_patterns.keys():
                    dates = matches.get(category, [])
                    row.append(", ".join(dates) if dates else "-")
                
                f.write("| " + " | ".join(row) + " |\n")
    
    def export_to_csv(self, data: List[Dict], output_path: str):
        """导出为CSV格式"""
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            # 写入表头
            header = ["主日期"] + list(self.category_patterns.keys())
            writer.writerow(header)
            
            # 写入数据
            for item in data:
                main_date = item["主日期"]
                matches = item["匹配"]
                
                row = [main_date]
                for category in self.category_patterns.keys():
                    dates = matches.get(category, [])
                    row.append(", ".join(dates) if dates else "")
                
                writer.writerow(row)
    
    def generate_statistics(self, data: List[Dict], output_path: str):
        """生成详细的统计信息"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("=== 增强版日期匹配统计信息 ===\n\n")
            f.write(f"总主日期数: {len(data)}\n\n")
            
            category_stats = {cat: 0 for cat in self.category_patterns.keys()}
            total_matches = 0
            
            for item in data:
                matches = item["匹配"]
                for category in self.category_patterns.keys():
                    if category in matches and matches[category]:
                        category_stats[category] += len(matches[category])
                        total_matches += len(matches[category])
            
            f.write("各分类匹配数量:\n")
            for category, count in category_stats.items():
                f.write(f"  {category}: {count} 个匹配\n")
            
            f.write(f"\n总匹配数: {total_matches}\n")
            
            # 计算平均匹配数
            if data:
                avg_matches = total_matches / len(data)
                f.write(f"平均每个主日期的匹配数: {avg_matches:.2f}\n")
            
            # 找出匹配最多的主日期
            max_matches = 0
            max_match_date = None
            for item in data:
                total_item_matches = sum(len(dates) for dates in item["匹配"].values())
                if total_item_matches > max_matches:
                    max_matches = total_item_matches
                    max_match_date = item["主日期"]
            
            if max_match_date:
                f.write(f"\n匹配最多的主日期: {max_match_date} ({max_matches} 个匹配)\n")
            
            # 双向匹配统计
            f.write("\n=== 双向匹配统计 ===\n")
            bidirectional_count = 0
            for item in data:
                main_date = item["主日期"]
                matches = item["匹配"]
                for category, dates in matches.items():
                    for date in dates:
                        # 检查是否存在反向匹配
                        if date in [d["主日期"] for d in data]:
                            target_item = next(d for d in data if d["主日期"] == date)
                            if category in target_item["匹配"] and main_date in target_item["匹配"][category]:
                                bidirectional_count += 1
            
            f.write(f"双向匹配数: {bidirectional_count}\n")
            f.write(f"双向匹配比例: {bidirectional_count/total_matches*100:.2f}%\n")

def main():
    """主函数"""
    print("=== 增强版日期匹配信息提取工具 ===\n")
    
    extractor = EnhancedDateExtractor('allday.docx')
    
    print("正在提取原始日期匹配信息...")
    data = extractor.extract_dates()
    
    print(f"成功提取 {len(data)} 个主日期的匹配信息")
    
    # 增强双向匹配关系
    enhanced_data = extractor.enhance_bidirectional_matches(data)
    
    # 保存增强版JSON数据
    with open('enhanced_date_matches.json', 'w', encoding='utf-8') as f:
        json.dump(enhanced_data, f, ensure_ascii=False, indent=2)
    print("✓ 增强版JSON数据已保存到 enhanced_date_matches.json")
    
    # 生成增强版反向索引
    reverse_index = extractor.generate_reverse_index(enhanced_data)
    with open('enhanced_reverse_index.json', 'w', encoding='utf-8') as f:
        json.dump(reverse_index, f, ensure_ascii=False, indent=2)
    print("✓ 增强版反向索引已保存到 enhanced_reverse_index.json")
    
    # 导出增强版Markdown表格
    extractor.export_to_markdown(enhanced_data, 'enhanced_date_matches.md')
    print("✓ 增强版Markdown表格已保存到 enhanced_date_matches.md")
    
    # 导出增强版CSV文件
    extractor.export_to_csv(enhanced_data, 'enhanced_date_matches.csv')
    print("✓ 增强版CSV文件已保存到 enhanced_date_matches.csv")
    
    # 生成增强版统计信息
    extractor.generate_statistics(enhanced_data, 'enhanced_statistics.txt')
    print("✓ 增强版统计信息已保存到 enhanced_statistics.txt")
    
    # 显示统计信息
    print("\n=== 增强版统计信息 ===")
    print(f"总主日期数: {len(enhanced_data)}")
    
    category_stats = {cat: 0 for cat in extractor.category_patterns.keys()}
    total_matches = 0
    
    for item in enhanced_data:
        matches = item["匹配"]
        for category in extractor.category_patterns.keys():
            if category in matches and matches[category]:
                category_stats[category] += len(matches[category])
                total_matches += len(matches[category])
    
    for category, count in category_stats.items():
        print(f"{category}: {count} 个匹配")
    
    print(f"总匹配数: {total_matches}")
    
    # 显示前几个示例
    print("\n=== 前3个增强版示例 ===")
    for i, item in enumerate(enhanced_data[:3]):
        print(f"\n主日期: {item['主日期']}")
        for category, dates in item['匹配'].items():
            if dates:
                print(f"  {category}: {', '.join(dates[:5])}{'...' if len(dates) > 5 else ''}")
    
    print("\n=== 增强版提取完成 ===")
    print("所有增强版文件已生成，现在所有匹配关系都是双向的！")

if __name__ == "__main__":
    main() 